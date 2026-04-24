#!/usr/bin/env python3
"""
选题推荐 Word 文档生成器（精美格式）
用法: echo '<json>' | python generate_report.py [--output path/to/file.docx]

输入 JSON 格式见 references/output_format.md 中的 report_payload 节。
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ── 颜色 ──────────────────────────────────────────────────────────────
C_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
C_WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
C_GRAY   = RGBColor(0x66, 0x66, 0x66)
C_LGRAY  = RGBColor(0xAA, 0xAA, 0xAA)
C_ORANGE = RGBColor(0xFF, 0x88, 0x00)

C_RED    = RGBColor(0xE5, 0x39, 0x35)   # 热点
C_GREEN  = RGBColor(0x2E, 0x7D, 0x32)   # 趋势
C_BLUE   = RGBColor(0x15, 0x65, 0xC0)   # 原创

TYPE_CFG = {
    "hot":      {"emoji": "🔥", "label": "热点蹭榜选题", "color": C_RED,   "bg_hdr": "C62828", "bg_card": "FFF3F3"},
    "trend":    {"emoji": "🌱", "label": "趋势借势选题", "color": C_GREEN, "bg_hdr": "1B5E20", "bg_card": "F1F8F1"},
    "original": {"emoji": "💡", "label": "原创高潜选题", "color": C_BLUE,  "bg_hdr": "0D47A1", "bg_card": "EFF4FF"},
}


# ── XML 辅助 ──────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex6: str):
    """设置单元格背景色，hex6 不含 #，例如 'FF4444'"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex6)
    tcPr.append(shd)


def _set_cell_border(cell, color_hex6: str = "DDDDDD"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        bd = OxmlElement(f"w:{edge}")
        bd.set(qn("w:val"), "single")
        bd.set(qn("w:sz"), "6")
        bd.set(qn("w:color"), color_hex6)
        borders.append(bd)
    tcPr.append(borders)


def _set_run_font(run, name: str = "微软雅黑"):
    run.font.name = name
    try:
        rPr = run._r.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.insert(0, rFonts)
        rFonts.set(qn("w:eastAsia"), name)
        rFonts.set(qn("w:ascii"), "Arial")
        rFonts.set(qn("w:hAnsi"), "Arial")
    except Exception:
        pass


def _run(para, text: str, bold=False, size=11,
         color: RGBColor | None = None, font="微软雅黑"):
    r = para.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    _set_run_font(r, font)
    return r


# ── 构建块 ────────────────────────────────────────────────────────────

def _divider(doc: Document, char="─", n=52, color: RGBColor = C_LGRAY):
    p = doc.add_paragraph(char * n)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.size = Pt(8)


def _vspace(doc: Document, pt: int = 4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(pt)


def _section_header(doc: Document, t: str):
    """带深色背景的分区标题条"""
    cfg = TYPE_CFG[t]
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, cfg["bg_hdr"])
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Cm(0.4)
    _run(p, f"{cfg['emoji']}  {cfg['label']}", bold=True, size=12, color=C_WHITE)
    _vspace(doc, 6)


def _rec_card(doc: Document, rec: dict, index: int):
    """单条选题卡片"""
    cfg = TYPE_CFG.get(rec.get("type", "hot"), TYPE_CFG["hot"])

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, cfg["bg_card"])
    _set_cell_border(cell, "DDDDDD")

    # ① 标题行
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after = Pt(3)
    p_title.paragraph_format.left_indent = Cm(0.4)
    _run(p_title, f"{index}.  ", bold=True, size=12, color=cfg["color"])
    _run(p_title, rec.get("title", ""), bold=True, size=12, color=C_DARK)

    # ② 热点标签（如有）
    hot = rec.get("hot_topic", "").strip()
    if hot:
        p_tag = cell.add_paragraph()
        p_tag.paragraph_format.space_before = Pt(0)
        p_tag.paragraph_format.space_after = Pt(3)
        p_tag.paragraph_format.left_indent = Cm(1.0)
        _run(p_tag, "📌 热点关联：", bold=True, size=9, color=C_GRAY)
        _run(p_tag, hot, size=9, color=cfg["color"])

    # ③ 匹配理由
    reason = rec.get("reason", "").strip()
    if reason:
        p_r = cell.add_paragraph()
        p_r.paragraph_format.space_before = Pt(0)
        p_r.paragraph_format.space_after = Pt(3)
        p_r.paragraph_format.left_indent = Cm(1.0)
        _run(p_r, "▸ 匹配理由：", bold=True, size=10, color=C_GRAY)
        _run(p_r, reason, size=10, color=C_DARK)

    # ④ 拍摄建议
    tips = rec.get("shooting_tips", "").strip()
    if tips:
        p_t = cell.add_paragraph()
        p_t.paragraph_format.space_before = Pt(0)
        p_t.paragraph_format.space_after = Pt(6)
        p_t.paragraph_format.left_indent = Cm(1.0)
        _run(p_t, "▸ 拍摄建议：", bold=True, size=10, color=C_GRAY)
        _run(p_t, tips, size=10, color=C_DARK)
    else:
        # 底部留白
        cell.paragraphs[-1].paragraph_format.space_after = Pt(6)

    _vspace(doc, 5)


def _trending_list(doc: Document, topics: list[dict]):
    """完整热榜列表，带序号和热度值"""
    if not topics:
        return

    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 表头
    for cell, label in zip(tbl.rows[0].cells, ["排名", "热点话题", "热度"]):
        _set_cell_bg(cell, "1A1A2E")
        _set_cell_border(cell, "1A1A2E")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        _run(p, label, bold=True, size=10, color=C_WHITE)

    for i, topic in enumerate(topics):
        row = tbl.add_row().cells
        bg = "F5F5F5" if i % 2 == 0 else "FFFFFF"
        for cell in row:
            _set_cell_bg(cell, bg)
            _set_cell_border(cell, "DDDDDD")

        # 排名
        p0 = row[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(2)
        p0.paragraph_format.space_after = Pt(2)
        rank_color = C_RED if i < 3 else C_DARK
        _run(p0, f"{i + 1}", bold=True, size=10, color=rank_color)

        # 热点话题
        p1 = row[1].paragraphs[0]
        p1.paragraph_format.left_indent = Cm(0.2)
        p1.paragraph_format.space_before = Pt(2)
        p1.paragraph_format.space_after = Pt(2)
        _run(p1, topic.get("title", ""), size=9, color=C_DARK)

        # 热度
        p2 = row[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_before = Pt(2)
        p2.paragraph_format.space_after = Pt(2)
        heat = topic.get("heat", 0)
        heat_str = f"{heat:,}" if heat else "—"
        _run(p2, heat_str, size=9, color=C_ORANGE)


def _summary_box(doc: Document, text: str):
    """今日方向建议黄色提示框"""
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    _set_cell_bg(cell, "FFF8E1")
    _set_cell_border(cell, "FFB300")

    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.4)
    _run(p, "💬  今日内容方向建议\n", bold=True, size=11,
         color=RGBColor(0xCC, 0x88, 0x00))
    _run(p, text, size=10, color=C_DARK)


# ── 主构建函数 ────────────────────────────────────────────────────────

def build_doc(data: dict, output_path: Path):
    doc = Document()

    # 页边距
    for sec in doc.sections:
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        sec.left_margin   = Cm(2.8)
        sec.right_margin  = Cm(2.8)

    prefs    = data.get("prefs", {})
    date_str = data.get("date", datetime.now().strftime("%Y-%m-%d %H:%M"))
    niche    = prefs.get("niche", "我的账号")
    audience = prefs.get("target_audience", "")
    topics   = data.get("topics", [])
    recs     = data.get("recommendations", [])
    daily    = data.get("daily_summary", "")

    # ── 大标题 ──────────────────────────────────────────
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after  = Pt(4)
    _run(p_title, "🎬  今日爆款选题推荐", bold=True, size=24, color=C_DARK)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(2)
    _run(p_sub, f"{date_str}  ·  ", size=11, color=C_GRAY)
    _run(p_sub, niche, bold=True, size=11, color=C_DARK)
    if audience:
        _run(p_sub, f"  ·  受众：{audience}", size=10, color=C_GRAY)

    _divider(doc)

    # ── 抖音热榜 TOP30 ───────────────────────────────────
    p_hdr = doc.add_paragraph()
    p_hdr.paragraph_format.space_after = Pt(6)
    _run(p_hdr, "📊  抖音热榜 TOP30", bold=True, size=13, color=C_DARK)

    _trending_list(doc, topics)
    _vspace(doc, 10)

    # ── 选题推荐（按分类分组）────────────────────────────
    grouped: dict[str, list] = {"hot": [], "trend": [], "original": []}
    for rec in recs:
        grouped.setdefault(rec.get("type", "hot"), []).append(rec)

    idx = 1
    for t in ("hot", "trend", "original"):
        group = grouped.get(t, [])
        if not group:
            continue
        _section_header(doc, t)
        for rec in group:
            _rec_card(doc, rec, idx)
            idx += 1

    # ── 今日方向总结 ─────────────────────────────────────
    if daily:
        _vspace(doc, 8)
        _summary_box(doc, daily)

    # ── 页脚 ─────────────────────────────────────────────
    _vspace(doc, 12)
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _run(p_foot, "由 Kimi · 抖音爆款选题 Skill 生成  ·  数据实时抓取，仅供参考",
         size=8, color=C_LGRAY)

    doc.save(str(output_path))
    print(f"✓ 文档已生成：{output_path}", file=sys.stderr)
    print(str(output_path))


# ── 入口 ─────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="生成选题推荐 Word 文档")
    parser.add_argument("--output", "-o", default="",
                        help="输出路径（默认保存到桌面）")
    args = parser.parse_args()

    raw = sys.stdin.read().strip()
    if not raw:
        print("错误：stdin 为空，请通过 stdin 传入 JSON 数据。", file=sys.stderr)
        sys.exit(1)

    try:
        data = json.loads(raw)
    except json.JSONDecodeError as e:
        print(f"JSON 解析失败: {e}", file=sys.stderr)
        sys.exit(1)

    if args.output:
        output_path = Path(args.output).expanduser()
    else:
        slug = datetime.now().strftime("%Y-%m-%d")
        output_path = Path.home() / "Desktop" / f"选题推荐_{slug}.docx"

    output_path.parent.mkdir(parents=True, exist_ok=True)
    build_doc(data, output_path)


if __name__ == "__main__":
    main()
